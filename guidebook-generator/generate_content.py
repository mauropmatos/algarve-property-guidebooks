#!/usr/bin/env python3
"""
Guidebook Content Generator
Produces a Word document (.docx) with all guidebook content in 5 languages,
ready to copy-paste into a Canva template.

Usage:
    python generate_content.py properties/casa_azul
    python generate_content.py            ← prompts for path interactively
"""

import sys
import argparse
import importlib.util
from pathlib import Path
from datetime import date

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: python-docx is not installed.")
    print("Run:   pip install python-docx")
    sys.exit(1)

# ── Section order and display names ───────────────────────────────────────────

SECTIONS = [
    ("welcome",          "Welcome Message"),
    ("checkin_checkout", "Check-in & Check-out"),
    ("house_rules",      "House Rules"),
    ("appliances",       "Appliances & Electricity"),
    ("local_area",       "Local Area Guide"),
    ("getting_around",   "Getting Around"),
    ("emergency",        "Emergency Contacts"),
    ("checkout",         "Checkout Checklist"),
]

LANGS = [
    ("en_uk", "English (UK)"),
    ("en_us", "English (US/CA)"),
    ("pt",    "Português"),
    ("fr",    "Français"),
    ("de",    "Deutsch"),
]

WIFI_LABELS = {
    "en_uk": ("Network", "Password"),
    "en_us": ("Network", "Password"),
    "pt":    ("Rede", "Palavra-passe"),
    "fr":    ("Réseau", "Mot de passe"),
    "de":    ("Netzwerk", "Passwort"),
}

# ── Config loader ──────────────────────────────────────────────────────────────

def load_config(prop_dir):
    spec = importlib.util.spec_from_file_location("config", prop_dir / "config.py")
    mod  = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod.PROPERTY

# ── Document helpers ───────────────────────────────────────────────────────────

def set_margins(doc, top=2.0, bottom=2.0, left=2.5, right=2.5):
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)


def add_title_block(doc, prop):
    """Document cover: property name + generated date."""
    title = doc.add_heading(prop["name"], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(
        f"Guidebook Content  ·  Generated {date.today().strftime('%d %B %Y')}  ·  "
        "Ready for Canva"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.runs[0]
    run.font.size  = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.italic = True

    doc.add_paragraph()  # breathing room


def add_section_heading(doc, number, title):
    """Heading 1 — section number + name."""
    h = doc.add_heading(f"{number}. {title}", level=1)
    return h


def add_lang_heading(doc, label):
    """Heading 2 — language label."""
    doc.add_heading(label, level=2)


def add_lang_divider(doc):
    """Thin horizontal rule between language blocks."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run("─" * 60)
    run.font.size  = Pt(7)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


def add_body_text(doc, text):
    """
    Render a plain-text block into Word paragraphs with lightweight formatting:
      ALL CAPS lines  → bold sub-heading
      • lines         → bullet list item
      □ / ✓ lines     → checkbox list item (List Bullet style)
      ⚠ lines         → shaded warning paragraph
      blank lines     → spacer
      everything else → Normal paragraph
    """
    for line in text.strip().splitlines():
        s = line.strip()

        if not s:
            doc.add_paragraph()
            continue

        if s.isupper() and len(s) > 3 and s[0] not in ("•", "□", "✓", "⚠"):
            p    = doc.add_paragraph()
            run  = p.add_run(s)
            run.font.bold = True
            run.font.size = Pt(10)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(2)

        elif s.startswith("•"):
            p = doc.add_paragraph(s[1:].strip(), style="List Bullet")
            p.paragraph_format.space_after = Pt(2)

        elif s.startswith("□") or s.startswith("✓"):
            p = doc.add_paragraph(s, style="List Bullet")
            p.paragraph_format.space_after = Pt(3)

        elif s.startswith("⚠"):
            p    = doc.add_paragraph()
            run  = p.add_run(s)
            run.font.bold  = True
            run.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)

        else:
            p = doc.add_paragraph(s)
            p.paragraph_format.space_after = Pt(3)


def add_wifi_table(doc, wifi, code):
    """Two-column table showing network name and password clearly."""
    net_lbl, pwd_lbl = WIFI_LABELS.get(code, ("Network", "Password"))

    tbl = doc.add_table(rows=2, cols=2)
    tbl.style = "Table Grid"

    # Header row
    for cell, text in zip(tbl.rows[0].cells, [net_lbl.upper(), pwd_lbl.upper()]):
        cell.text = text
        run = cell.paragraphs[0].runs[0]
        run.font.bold  = True
        run.font.size  = Pt(8)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Value row
    for cell, text in zip(tbl.rows[1].cells,
                           [wifi.get("network", "—"), wifi.get("password", "—")]):
        cell.text = text
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(15)

    doc.add_paragraph()  # space after table


# ── Document builder ───────────────────────────────────────────────────────────

def generate_doc(prop, output_path):
    doc = Document()
    set_margins(doc)

    # ── Cover ─────────────────────────────────────────────────────────────────
    add_title_block(doc, prop)

    # ── 1. Property Name & Tagline ────────────────────────────────────────────
    add_section_heading(doc, 1, "Property Name & Tagline")

    tagline = prop.get("tagline", {})
    for code, label in LANGS:
        add_lang_heading(doc, label)

        # Property name (same in every language)
        p = doc.add_paragraph()
        r = p.add_run(prop["name"])
        r.font.bold = True
        r.font.size = Pt(16)

        # Tagline (translated)
        tag = tagline.get(code, "") if isinstance(tagline, dict) else str(tagline)
        if tag:
            tp = doc.add_paragraph(tag)
            tp.runs[0].font.italic = True
            tp.runs[0].font.size   = Pt(12)

        add_lang_divider(doc)

    doc.add_page_break()

    # ── 2. WiFi ───────────────────────────────────────────────────────────────
    add_section_heading(doc, 2, "WiFi")

    wifi = prop.get("wifi", {})
    for code, label in LANGS:
        add_lang_heading(doc, label)
        add_wifi_table(doc, wifi, code)
        add_lang_divider(doc)

    doc.add_page_break()

    # ── 3–10. Content sections ────────────────────────────────────────────────
    content = prop.get("content", {})

    for i, (key, title) in enumerate(SECTIONS, start=3):
        add_section_heading(doc, i, title)

        section_data = content.get(key, {})
        for code, label in LANGS:
            text = section_data.get(code, "").strip()
            if not text:
                continue
            add_lang_heading(doc, label)
            add_body_text(doc, text)
            add_lang_divider(doc)

        doc.add_page_break()

    doc.save(str(output_path))
    print(f"  ✓  {output_path}")


# ── Entry point ────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Generate a Word guidebook content file from a property config."
    )
    parser.add_argument(
        "property_dir", nargs="?",
        help="Path to the property folder containing config.py"
    )
    args = parser.parse_args()

    if args.property_dir:
        prop_dir = Path(args.property_dir).resolve()
    else:
        raw = input("Property folder path: ").strip().strip('"').strip("'")
        prop_dir = Path(raw).resolve()

    if not (prop_dir / "config.py").exists():
        print(f"Error: config.py not found in {prop_dir}")
        sys.exit(1)

    prop = load_config(prop_dir)
    slug = prop["name"].lower().replace(" ", "-").replace("/", "-")

    output_dir = Path(__file__).parent / "output"
    output_dir.mkdir(exist_ok=True)
    output_path = output_dir / f"{slug}-content.docx"

    print(f"\nGenerating: {prop['name']}")
    generate_doc(prop, output_path)
    print("Done.\n")


if __name__ == "__main__":
    main()
